import os, base64, json, re, itertools
import fitz
import requests
from PIL import Image
from io import BytesIO
from docx import Document
from fastapi import FastAPI, Header, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from google import genai
from groq import Groq
from dotenv import load_dotenv

load_dotenv()

app = FastAPI(title="DocMind API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# ---------------- KEYS ----------------
MY_API_KEY = os.getenv("MY_API_KEY", "sk_docmind_2026_secure")
VISION_API_KEY = os.getenv("VISION_API_KEY", "")
GROQ_KEY = os.getenv("GROQ_KEY")

GEMINI_KEYS = [k for k in [
    os.getenv("GEMINI_KEY_1"),
    os.getenv("GEMINI_KEY_2"),
    os.getenv("GEMINI_KEY_3"),
    os.getenv("GEMINI_KEY_4"),
] if k]

_cycle = itertools.cycle(GEMINI_KEYS) if GEMINI_KEYS else None

def next_client():
    if not _cycle:
        raise Exception("No Gemini keys configured")
    return genai.Client(api_key=next(_cycle))


# ---------------- REQUEST MODEL ----------------
class DocRequest(BaseModel):
    fileName: str
    fileType: str
    fileBase64: str


# ---------------- TEXT EXTRACTION ----------------
def extract_pdf(b):
    try:
        doc = fitz.open(stream=b, filetype="pdf")
        text = "\n".join(page.get_text() for page in doc).strip()
        if len(text) > 20:
            return text
    except:
        pass
    return ""


def extract_docx(b):
    try:
        doc = Document(BytesIO(b))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        text = "\n".join(parts).strip()
        if len(text) > 20:
            return text
    except:
        pass
    return ""


def extract_image_vision(b):
    if not VISION_API_KEY:
        return ""
    try:
        img = Image.open(BytesIO(b))
        if img.mode != "RGB":
            img = img.convert("RGB")

        if max(img.size) > 2000:
            img.thumbnail((2000, 2000), Image.LANCZOS)

        buf = BytesIO()
        img.save(buf, format="JPEG", quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode()

        url = f"https://vision.googleapis.com/v1/images:annotate?key={VISION_API_KEY}"

        response = requests.post(
            url,
            json={
                "requests": [{
                    "image": {"content": b64},
                    "features": [{"type": "TEXT_DETECTION", "maxResults": 1}]
                }]
            },
            timeout=30
        )

        data = response.json()
        annotations = data.get("responses", [{}])[0].get("textAnnotations", [])
        if annotations:
            return annotations[0].get("description", "").strip()

    except:
        pass

    return ""


def extract_image_gemini(b):
    try:
        img = Image.open(BytesIO(b))
        if img.mode != "RGB":
            img = img.convert("RGB")

        buf = BytesIO()
        img.save(buf, format="JPEG", quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode()

        client = next_client()

        r = client.models.generate_content(
            model="gemini-2.0-flash",
            contents=[{
                "role": "user",
                "parts": [
                    {"text": "Extract ALL text from this image. Return only raw text."},
                    {"inline_data": {"mime_type": "image/jpeg", "data": b64}}
                ]
            }]
        )

        raw = r.text
        if not raw and hasattr(r, "candidates") and r.candidates:
            raw = r.candidates[0].content.parts[0].text

        return raw.strip() if raw else ""

    except:
        return ""


def extract_text(b, file_type):
    ft = file_type.lower().strip()

    if ft == "pdf":
        text = extract_pdf(b)
        if len(text) > 20:
            return text

        try:
            doc = fitz.open(stream=b, filetype="pdf")
            parts = []
            for page in doc:
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("jpeg")
                t = extract_image_vision(img_bytes) or extract_image_gemini(img_bytes)
                if t:
                    parts.append(t)
            if parts:
                return "\n".join(parts)
        except:
            pass

    if ft in ("docx", "doc"):
        text = extract_docx(b)
        if len(text) > 20:
            return text

    if ft in ("image", "jpg", "jpeg", "png", "gif", "bmp", "tiff", "webp"):
        text = extract_image_vision(b)
        if len(text) > 20:
            return text
        return extract_image_gemini(b)

    return ""


# ---------------- AI ANALYSIS ----------------
PROMPT = """Analyze the document text. Return ONLY raw JSON, no markdown.

Document:
\"\"\"{text}\"\"\"

Return exactly this JSON:
{{
  "summary": "2-3 sentence accurate summary",
  "entities": {{
    "names": [],
    "dates": [],
    "organizations": [],
    "amounts": []
  }},
  "sentiment": "Positive or Neutral or Negative"
}}
"""


def clean_json(raw):
    raw = re.sub(r"^```json\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"^```\s*", "", raw, flags=re.MULTILINE)
    raw = re.sub(r"\s*```$", "", raw, flags=re.MULTILINE)
    return raw.strip()


def analyze(text):
    prompt = PROMPT.format(text=text[:6000])

    # -------- GEMINI --------
    for _ in range(len(GEMINI_KEYS)):
        try:
            client = next_client()
            r = client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt
            )

            raw = r.text
            if not raw and hasattr(r, "candidates") and r.candidates:
                raw = r.candidates[0].content.parts[0].text

            if raw:
                raw = clean_json(raw)
                result = json.loads(raw)
                if "summary" in result and "entities" in result:
                    return result

        except:
            continue

    # -------- GROQ FALLBACK --------
    if GROQ_KEY:
        try:
            groq_client = Groq(api_key=GROQ_KEY)
            r = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=1000
            )

            raw = r.choices[0].message.content
            if raw:
                raw = clean_json(raw)
                result = json.loads(raw)
                if "summary" in result and "entities" in result:
                    return result

        except:
            pass

    return {
        "summary": "Document processed.",
        "entities": {"names": [], "dates": [], "organizations": [], "amounts": []},
        "sentiment": "Neutral"
    }


# ---------------- API ENDPOINT ----------------
@app.post("/api/document-analyze")
async def analyze_document(req: DocRequest, x_api_key: str = Header(None)):
    if not x_api_key or x_api_key != MY_API_KEY:
        raise HTTPException(status_code=401, detail="Unauthorized")

    try:
        file_bytes = base64.b64decode(req.fileBase64)
    except:
        raise HTTPException(status_code=400, detail="Invalid base64")

    text = extract_text(file_bytes, req.fileType)

    if not text or len(text.strip()) < 5:
        return {
            "status": "success",
            "fileName": req.fileName,
            "summary": "Document received but contains no extractable text.",
            "entities": {"names": [], "dates": [], "organizations": [], "amounts": []},
            "sentiment": "Neutral"
        }

    result = analyze(text)

    return {
        "status": "success",
        "fileName": req.fileName,
        "summary": result.get("summary", ""),
        "entities": {
            "names": result.get("entities", {}).get("names", []),
            "dates": result.get("entities", {}).get("dates", []),
            "organizations": result.get("entities", {}).get("organizations", []),
            "amounts": result.get("entities", {}).get("amounts", [])
        },
        "sentiment": result.get("sentiment", "Neutral")
    }


@app.get("/")
def root():
    return {"status": "DocMind API is running", "version": "2.0.0"}


@app.get("/health")
@app.head("/health")
def health():
    return {"status": "healthy"}